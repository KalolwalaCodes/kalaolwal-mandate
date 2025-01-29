import sys
import json
from docx import Document
import re

def replace_text_with_regex(doc, old_texts, new_texts):
    """
    Replace text using regex in paragraphs, tables, headers, footers, and textboxes.
    Preserves styles while replacing text.
    """
    # Normalize text by removing non-breaking spaces and other hidden characters
    def normalize_text(text):
        return text.replace("\xa0", " ").strip()  # Replace non-breaking spaces with normal spaces

    # Function to replace text and preserve styles
    def replace_with_style(paragraph, old_text, new_text):
        for run in paragraph.runs:
            if old_text in run.text:
                # Log the match for debugging
                print(f"Found match: {old_text} in paragraph run: {run.text}")
                # Preserve the original style
                run.text = run.text.replace(old_text, new_text)

    # Iterate through all pairs of old and new text
    for old_text, new_text in zip(old_texts, new_texts):
        pattern = re.escape(old_text)  # Escape the old text for regex

        # Replace text in paragraphs
        for paragraph in doc.paragraphs:
            normalized_text = normalize_text(paragraph.text)
            if re.search(pattern, normalized_text):
                print(f"Replacing in paragraph: {paragraph.text}")
                replace_with_style(paragraph, old_text, new_text)

        # Replace text in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    normalized_text = normalize_text(cell.text)
                    if re.search(pattern, normalized_text):
                        print(f"Replacing in cell: {cell.text}")
                        replace_with_style(cell, old_text, new_text)

        # Check headers and footers for text
        for section in doc.sections:
            header = section.header
            for paragraph in header.paragraphs:
                normalized_text = normalize_text(paragraph.text)
                if re.search(pattern, normalized_text):
                    print(f"Replacing in header: {paragraph.text}")
                    replace_with_style(paragraph, old_text, new_text)

            footer = section.footer
            for paragraph in footer.paragraphs:
                normalized_text = normalize_text(paragraph.text)
                if re.search(pattern, normalized_text):
                    print(f"Replacing in footer: {paragraph.text}")
                    replace_with_style(paragraph, old_text, new_text)

def replace_text_in_docx(input_path, output_path, old_texts, new_texts):
    """
    Main function to handle .docx text replacement using regex.
    """
    doc = Document(input_path)

    replace_text_with_regex(doc, old_texts, new_texts)

    # Save the updated document
    doc.save(output_path)
    print(f"Updated file saved at: {output_path}")

if __name__ == "__main__":
    input_path = sys.argv[1]
    output_path = sys.argv[2]

    # Parse the JSON arrays passed from Node.js
    old_texts = json.loads(sys.argv[3])
    new_texts = json.loads(sys.argv[4])

    replace_text_in_docx(input_path, output_path, old_texts, new_texts)
